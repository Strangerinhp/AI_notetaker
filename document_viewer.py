"""Safe, dependency-light HTML preview for stored DOCX meeting reports."""

from __future__ import annotations

import html
import io
import zipfile
from collections.abc import Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


MAX_UNCOMPRESSED_BYTES = 120 * 1024 * 1024
MAX_ARCHIVE_MEMBERS = 5_000


class InvalidWordDocument(ValueError):
    """Raised when an uploaded file is not a safe, readable DOCX document."""


def validate_docx(data: bytes) -> None:
    """Reject malformed files and simple ZIP bombs before python-docx parses them."""
    if not data or not data.startswith(b"PK"):
        raise InvalidWordDocument("File đã chọn không phải tài liệu DOCX hợp lệ.")
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            names = {member.filename for member in members}
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise InvalidWordDocument("File đã chọn không phải tài liệu DOCX hợp lệ.")
            if len(members) > MAX_ARCHIVE_MEMBERS:
                raise InvalidWordDocument("Tài liệu Word chứa quá nhiều thành phần.")
            if sum(member.file_size for member in members) > MAX_UNCOMPRESSED_BYTES:
                raise InvalidWordDocument("Tài liệu Word quá lớn sau khi giải nén.")
            if any(member.flag_bits & 0x1 for member in members):
                raise InvalidWordDocument("Tài liệu Word được mã hóa không được hỗ trợ.")
    except (zipfile.BadZipFile, OSError) as error:
        raise InvalidWordDocument("Không đọc được cấu trúc tài liệu DOCX.") from error

    try:
        Document(io.BytesIO(data))
    except Exception as error:
        raise InvalidWordDocument("Microsoft Word có thể mở file này, nhưng viewer không đọc được.") from error


def _iter_blocks(parent: DocumentObject | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocumentObject):
        container = parent.element.body
    else:
        container = parent._tc
    for child in container.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _length_points(value) -> float | None:
    try:
        return round(float(value.pt), 2) if value is not None else None
    except (AttributeError, TypeError, ValueError):
        return None


def _paragraph_style(paragraph: Paragraph) -> str:
    fmt = paragraph.paragraph_format
    rules = []
    alignment = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    }.get(paragraph.alignment)
    if alignment:
        rules.append(f"text-align:{alignment}")
    for attribute, css_name in (
        (fmt.left_indent, "margin-left"),
        (fmt.right_indent, "margin-right"),
        (fmt.first_line_indent, "text-indent"),
        (fmt.space_before, "margin-top"),
        (fmt.space_after, "margin-bottom"),
    ):
        points = _length_points(attribute)
        if points is not None:
            rules.append(f"{css_name}:{points}pt")
    line_spacing = fmt.line_spacing
    if isinstance(line_spacing, (int, float)):
        rules.append(f"line-height:{max(float(line_spacing), 0.8):.2f}")
    else:
        points = _length_points(line_spacing)
        if points:
            rules.append(f"line-height:{points}pt")
    return ";".join(rules)


def _render_run(run, paragraph: Paragraph) -> str:
    text = html.escape(run.text).replace("\t", "&emsp;").replace("\n", "<br>")
    if not text:
        return ""
    paragraph_font = paragraph.style.font if paragraph.style else None
    rules = []
    font_name = run.font.name or (paragraph_font.name if paragraph_font else None)
    font_size = run.font.size or (paragraph_font.size if paragraph_font else None)
    if font_name:
        safe_font_name = "".join(
            character
            for character in font_name[:80]
            if character.isalnum() or character in " -_,."
        )
        if safe_font_name:
            rules.append(f'font-family:"{safe_font_name}",serif')
    points = _length_points(font_size)
    if points:
        rules.append(f"font-size:{points}pt")
    if run.font.color and run.font.color.rgb:
        rules.append(f"color:#{run.font.color.rgb}")
    style = f' style="{";".join(rules)}"' if rules else ""
    content = f"<span{style}>{text}</span>"
    bold = run.bold if run.bold is not None else bool(paragraph_font and paragraph_font.bold)
    italic = run.italic if run.italic is not None else bool(paragraph_font and paragraph_font.italic)
    if run.underline:
        content = f"<u>{content}</u>"
    if italic:
        content = f"<em>{content}</em>"
    if bold:
        content = f"<strong>{content}</strong>"
    return content


def _numbering_marker(paragraph: Paragraph) -> str:
    properties = paragraph._p.pPr
    numbered = properties is not None and properties.numPr is not None
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if not numbered and "list" not in style_name:
        return ""
    marker = "1." if "number" in style_name else "•"
    return f'<span class="list-marker">{marker}</span>'


def _render_paragraph(paragraph: Paragraph) -> str:
    style_name = paragraph.style.name if paragraph.style else ""
    lowered = style_name.lower()
    tag = "p"
    for level in range(1, 7):
        if f"heading {level}" in lowered or f"tiêu đề {level}" in lowered:
            tag = f"h{level}"
            break
    content = "".join(_render_run(run, paragraph) for run in paragraph.runs)
    if not content and paragraph.text:
        content = html.escape(paragraph.text)
    marker = _numbering_marker(paragraph)
    if not content:
        content = "&nbsp;"
    inline_style = _paragraph_style(paragraph)
    style = f' style="{inline_style}"' if inline_style else ""
    css_class = " has-marker" if marker else ""
    return f'<{tag} class="word-paragraph{css_class}"{style}>{marker}{content}</{tag}>'


def _table_is_borderless(table: Table) -> bool:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders")) if properties is not None else None
    if borders is None:
        return False
    edges = list(borders)
    return bool(edges) and all(
        edge.get(qn("w:val"), "").lower() in {"nil", "none"}
        for edge in edges
    )


def _render_cell(cell: _Cell) -> str:
    blocks = []
    for block in _iter_blocks(cell):
        blocks.append(
            _render_paragraph(block)
            if isinstance(block, Paragraph)
            else _render_table(block)
        )
    return "".join(blocks) or "&nbsp;"


def _render_table(table: Table) -> str:
    css_class = "word-table borderless" if _table_is_borderless(table) else "word-table"
    rows = []
    for row in table.rows:
        rendered_cells = []
        seen_cells = set()
        for cell in row.cells:
            identity = id(cell._tc)
            if identity in seen_cells:
                continue
            seen_cells.add(identity)
            grid_span = cell._tc.tcPr.gridSpan
            span = int(grid_span.val) if grid_span is not None else 1
            colspan = f' colspan="{span}"' if span > 1 else ""
            rendered_cells.append(f"<td{colspan}>{_render_cell(cell)}</td>")
        rows.append(f"<tr>{''.join(rendered_cells)}</tr>")
    return f'<table class="{css_class}"><tbody>{"".join(rows)}</tbody></table>'


def render_docx_html(data: bytes, *, title: str = "Báo cáo cuộc họp") -> str:
    """Render the main DOCX body as a read-only, Word-like HTML page."""
    validate_docx(data)
    document = Document(io.BytesIO(data))
    section = document.sections[0]
    page_width = round(section.page_width.mm, 2)
    page_height = round(section.page_height.mm, 2)
    top = round(section.top_margin.mm, 2)
    right = round(section.right_margin.mm, 2)
    bottom = round(section.bottom_margin.mm, 2)
    left = round(section.left_margin.mm, 2)
    body = "".join(
        _render_paragraph(block)
        if isinstance(block, Paragraph)
        else _render_table(block)
        for block in _iter_blocks(document)
    )
    safe_title = html.escape(title)
    return f"""<!doctype html>
<html lang="vi">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>{safe_title}</title>
  <style>
    * {{ box-sizing: border-box; }}
    html {{ background: #dfe1e5; }}
    body {{ margin: 0; padding: 28px; color: #111; background: #dfe1e5; }}
    .word-page {{ width: {page_width}mm; min-height: {page_height}mm; margin: 0 auto;
      padding: {top}mm {right}mm {bottom}mm {left}mm; background: #fff;
      box-shadow: 0 2px 14px rgba(20,25,35,.18); font-family: "Times New Roman",serif;
      font-size: 13pt; line-height: 1.25; }}
    .word-paragraph {{ min-height: 1em; margin: 0 0 6pt; white-space: pre-wrap;
      overflow-wrap: anywhere; }}
    h1.word-paragraph, h2.word-paragraph, h3.word-paragraph {{ font-weight: 700; }}
    .has-marker {{ position: relative; padding-left: 18pt; }}
    .list-marker {{ position: absolute; left: 0; min-width: 14pt; }}
    .word-table {{ width: 100%; margin: 0 0 6pt; border-collapse: collapse;
      table-layout: fixed; }}
    .word-table td {{ padding: 2pt 4pt; vertical-align: top; border: 1px solid #777; }}
    .word-table.borderless td {{ border: 0; }}
    .word-table .word-paragraph {{ margin-bottom: 0; }}
    @media (max-width: 900px) {{
      body {{ padding: 14px; }}
      .word-page {{ transform-origin: top left; }}
    }}
    @media print {{ body {{ padding: 0; background: #fff; }} .word-page {{ box-shadow: none; }} }}
  </style>
</head>
<body><main class="word-page">{body}</main></body>
</html>"""
