"""Export editable meeting Markdown as a PTC1-style Word report."""

from __future__ import annotations

import re
import shutil
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips
from lxml import etree

BASE_DIR = Path(__file__).resolve().parent
REPORT_TEMPLATE_PATH = BASE_DIR / "templates" / "ptc1_report_template.docx"
REPORT_TITLE_PLACEHOLDER = "[TÊN BÁO CÁO]"
REPORT_CONTENT_PLACEHOLDER = "[NỘI DUNG BIÊN BẢN ĐƯỢC CHÈN TẠI ĐÂY]"

FONT_NAME = "Times New Roman"
BODY_SIZE = 14
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
CONTENT_WIDTH_CM = 16.0
HEADER_TABLE_WIDTHS_DXA = (4321, 5779)
SIGNOFF_TABLE_WIDTHS_DXA = (4619, 4452)

WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
LIST_RE = re.compile(r"^(\s*)(?:[-+*]|\d+[.)])\s+(.+?)\s*$")
TABLE_RULE_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$")
METADATA_RE = re.compile(
    r"^\*\*(?:Thời gian|Địa điểm|Chủ trì|Thành phần tham dự)\s*:\*\*",
    re.IGNORECASE,
)
INLINE_RE = re.compile(
    r"\*\*\*(.+?)\*\*\*|"
    r"\*\*(.+?)\*\*|"
    r"__(.+?)__|"
    r"(?<!\*)\*([^*\n]+?)\*|"
    r"(?<!_)_([^_\n]+?)_|"
    r"`([^`\n]+)`|"
    r"\[([^]\n]+)\]\(([^)\n]+)\)"
)
REPORT_DATE_RE = re.compile(
    r"Hà Nội,\s*ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}",
    re.IGNORECASE,
)


class ReportTemplateError(ValueError):
    """Raised when the retained report template no longer has its required slots."""


def _set_run_font(
    run,
    *,
    size: float = BODY_SIZE,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    for font_key in ("ascii", "hAnsi", "eastAsia", "cs"):
        run_fonts.set(qn(f"w:{font_key}"), FONT_NAME)


def _clear_story_container(container) -> None:
    root = container._element
    for child in list(root):
        root.remove(child)
    root.append(OxmlElement("w:p"))


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, *, top=0, start=0, bottom=0, end=0) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry_dxa(
    table,
    widths_dxa: tuple[int, ...] | list[int],
    *,
    alignment=WD_TABLE_ALIGNMENT.CENTER,
    zero_cell_margins: bool = True,
) -> None:
    table.alignment = alignment
    table.autofit = False
    total_width = sum(widths_dxa)
    properties = table._tbl.tblPr

    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total_width))
    table_width.set(qn("w:type"), "dxa")

    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_columns = table._tbl.tblGrid.gridCol_lst
    for index, width_dxa in enumerate(widths_dxa):
        width = Twips(width_dxa)
        if index < len(grid_columns):
            grid_columns[index].set(qn("w:w"), str(width_dxa))
        for row in table.rows:
            cell = row.cells[index]
            cell.width = width
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(width_dxa))
            cell_width.set(qn("w:type"), "dxa")
            if zero_cell_margins:
                _set_cell_margins(cell)


def _set_table_geometry(table, widths_cm: list[float]) -> None:
    _set_table_geometry_dxa(
        table,
        [Cm(width).twips for width in widths_cm],
    )


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    _set_run_font(run, size=12)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, display, end))


def _configure_document(document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.0)
    section.footer_distance = Cm(0.0)
    section.different_first_page_header_footer = True

    for header in (
        section.header,
        section.first_page_header,
        section.even_page_header,
    ):
        _clear_story_container(header)
    _add_page_number(section.header.paragraphs[0])
    for footer in (
        section.footer,
        section.first_page_footer,
        section.even_page_footer,
    ):
        _clear_story_container(footer)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_SIZE)
    normal_rpr = normal._element.get_or_add_rPr()
    normal_fonts = normal_rpr.get_or_add_rFonts()
    for font_key in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{font_key}"), FONT_NAME)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(3)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    for level in range(1, 4):
        style_name = f"Meeting Heading {level}"
        if style_name in document.styles:
            style = document.styles[style_name]
        else:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style.font.size = Pt(BODY_SIZE)
        style.font.bold = True
        style_rpr = style._element.get_or_add_rPr()
        style_fonts = style_rpr.get_or_add_rFonts()
        for font_key in ("ascii", "hAnsi", "eastAsia", "cs"):
            style_fonts.set(qn(f"w:{font_key}"), FONT_NAME)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0


def _add_inline_runs(paragraph, text: str, *, size: float = BODY_SIZE) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            _set_run_font(paragraph.add_run(text[cursor : match.start()]), size=size)

        (
            bold_italic_text,
            bold_text,
            bold_underscore_text,
            italic_text,
            italic_underscore_text,
            code_text,
            link_text,
            link_url,
        ) = match.groups()
        if bold_italic_text is not None:
            _set_run_font(
                paragraph.add_run(bold_italic_text),
                size=size,
                bold=True,
                italic=True,
            )
        elif bold_text is not None or bold_underscore_text is not None:
            _set_run_font(
                paragraph.add_run(bold_text or bold_underscore_text),
                size=size,
                bold=True,
            )
        elif italic_text is not None or italic_underscore_text is not None:
            _set_run_font(
                paragraph.add_run(italic_text or italic_underscore_text),
                size=size,
                italic=True,
            )
        elif code_text is not None:
            run = paragraph.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(max(10, size - 2))
        else:
            label = link_text or link_url
            value = label if label == link_url else f"{label} ({link_url})"
            run = paragraph.add_run(value)
            _set_run_font(run, size=size)
            run.underline = True
        cursor = match.end()

    if cursor < len(text):
        _set_run_font(paragraph.add_run(text[cursor:]), size=size)


def _format_body_paragraph(paragraph, *, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.first_line_indent = Cm(1.0 if first_line else 0)


def _add_heading(document, text: str, level: int, *, title_slot: bool = False) -> None:
    paragraph = document.add_paragraph(style=f"Meeting Heading {min(level, 3)}")
    if level == 1 or title_slot:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0 if level == 1 else 6)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_before = Pt(3)
    _add_inline_runs(paragraph, text, size=BODY_SIZE)
    for run in paragraph.runs:
        run.bold = True


def _add_rule(document) -> None:
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    properties.append(borders)


def _split_table_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", value)]


def _add_markdown_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    equal_width = CONTENT_WIDTH_CM / column_count
    _set_table_geometry(table, [equal_width] * column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            _set_cell_margins(cell, top=90, start=100, bottom=90, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = values[column_index] if column_index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            _add_inline_runs(paragraph, text, size=12)
            if row_index == 0:
                _set_cell_shading(cell, "EDEDED")
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def _add_markdown(document, markdown: str) -> None:
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    paragraph_buffer: list[str] = []
    index = 0
    section_started = False

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer if part.strip())
        paragraph_buffer.clear()
        if not text:
            return
        paragraph = document.add_paragraph()
        _format_body_paragraph(
            paragraph,
            first_line=not section_started and not METADATA_RE.match(text),
        )
        _add_inline_runs(paragraph, text)

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            index += 1
            continue

        heading = HEADING_RE.match(stripped)
        if heading:
            flush_paragraph()
            section_started = True
            _add_heading(
                document,
                heading.group(2).strip(" *_`"),
                len(heading.group(1)),
            )
            index += 1
            continue

        if stripped in {"---", "***", "___"}:
            flush_paragraph()
            _add_rule(document)
            index += 1
            continue

        if index + 1 < len(lines) and "|" in stripped and TABLE_RULE_RE.match(lines[index + 1]):
            flush_paragraph()
            rows = [_split_table_row(stripped)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_split_table_row(lines[index]))
                index += 1
            _add_markdown_table(document, rows)
            continue

        list_item = LIST_RE.match(line)
        if list_item:
            flush_paragraph()
            indent_level = min(4, len(list_item.group(1).replace("\t", "    ")) // 2)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.7)
            paragraph.paragraph_format.first_line_indent = Cm(-0.45)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            _set_run_font(paragraph.add_run("- "))
            _add_inline_runs(paragraph, list_item.group(2))
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1.0)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            _add_inline_runs(paragraph, stripped.lstrip("> "), size=13)
            for run in paragraph.runs:
                run.italic = True
            index += 1
            continue

        if METADATA_RE.match(stripped):
            flush_paragraph()
            paragraph = document.add_paragraph()
            _format_body_paragraph(paragraph, first_line=False)
            _add_inline_runs(paragraph, stripped)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def _strip_formal_markdown_title(markdown: str) -> str:
    lines = markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    first = next((index for index, line in enumerate(lines) if line.strip()), None)
    if first is None:
        return ""
    match = HEADING_RE.match(lines[first].strip())
    if not match or len(match.group(1)) != 1 or match.group(2).strip().upper() != "THÔNG BÁO":
        return markdown

    del lines[first]
    second = next((index for index, line in enumerate(lines) if line.strip()), None)
    if second is not None:
        second_match = HEADING_RE.match(lines[second].strip())
        if second_match and len(second_match.group(1)) == 2:
            del lines[second]
    return "\n".join(lines).lstrip("\n")


def _add_transcript(document, markdown: str, title: str) -> None:
    _add_heading(document, "TRANSCRIPT CUỘC HỌP", 1, title_slot=True)
    if title:
        _add_heading(document, title, 2, title_slot=True)
    for raw_line in markdown.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        paragraph = document.add_paragraph()
        _format_body_paragraph(paragraph, first_line=False)
        _add_inline_runs(paragraph, line, size=12)


def _new_transcript_document():
    document = Document()
    _configure_document(document)
    return document


def _load_report_template(template_path: Path | None = None):
    source = Path(template_path) if template_path else REPORT_TEMPLATE_PATH
    if not source.is_file():
        raise ReportTemplateError(f"Không tìm thấy template báo cáo Word: {source}")
    return Document(source)


def _replace_slot_paragraph(paragraph, value: str) -> None:
    """Replace a whole-paragraph slot while retaining its original formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _find_body_slot(document, placeholder: str):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == placeholder
    ]
    if len(matches) != 1:
        raise ReportTemplateError(
            f"Template Word phải có đúng một đoạn chứa {placeholder}."
        )
    return matches[0]


def _replace_report_date(document, report_date: date) -> None:
    replacement = (
        f"Hà Nội, ngày {report_date.day:02d} tháng {report_date.month:02d} "
        f"năm {report_date.year}"
    )
    matches = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if REPORT_DATE_RE.fullmatch(paragraph.text.strip()):
                        matches.append(paragraph)
    if len(matches) != 1:
        raise ReportTemplateError(
            "Template Word phải có đúng một dòng ngày dạng "
            "'Hà Nội, ngày DD tháng MM năm YYYY'."
        )
    _replace_slot_paragraph(matches[0], replacement)


def _insert_markdown_at_report_slot(document, markdown: str) -> None:
    marker = _find_body_slot(document, REPORT_CONTENT_PLACEHOLDER)
    body = document._element.body
    existing_nodes = set(body.iterchildren())

    _add_markdown(document, markdown)
    generated_nodes = [
        node
        for node in body.iterchildren()
        if node not in existing_nodes and node.tag != qn("w:sectPr")
    ]
    for node in generated_nodes:
        marker._p.addprevious(node)
    marker._p.getparent().remove(marker._p)


def _fill_report_template(
    document,
    markdown: str,
    title: str,
    report_date: date,
) -> None:
    title_slot = _find_body_slot(document, REPORT_TITLE_PLACEHOLDER)
    _replace_slot_paragraph(title_slot, title)
    _replace_report_date(document, report_date)
    _insert_markdown_at_report_slot(
        document,
        _strip_formal_markdown_title(markdown),
    )


def _set_clean_properties(document, title: str) -> None:
    properties = document.core_properties
    properties.title = title
    properties.subject = "Thông báo kết luận cuộc họp - PTC1"
    properties.author = "Công ty Truyền tải điện 1"
    properties.last_modified_by = "MeetNote"
    properties.comments = ""
    properties.keywords = ""


def export_markdown_to_docx(
    markdown: str,
    output: BinaryIO | BytesIO,
    *,
    title: str = "",
    document_type: str = "summary",
    report_date: date | datetime | None = None,
    template_path: Path | None = None,
) -> None:
    """Write summary/transcript Markdown as an editable `.docx` file."""
    normalized_title = title.strip() or "Kết luận cuộc họp"

    if document_type == "transcript":
        document = _new_transcript_document()
        _set_clean_properties(document, normalized_title)
        _add_transcript(document, markdown, normalized_title)
    else:
        document = _load_report_template(template_path)
        _set_clean_properties(document, normalized_title)
        effective_date = (
            report_date.date()
            if isinstance(report_date, datetime)
            else report_date
        )
        _fill_report_template(
            document,
            markdown,
            normalized_title,
            effective_date or datetime.now().date(),
        )

    document.save(output)


def create_reusable_template(reference_path: str | Path, output_path: str | Path) -> None:
    """Validate and copy a retained PTC1 template without rebuilding its layout."""
    reference = Path(reference_path)
    output = Path(output_path)
    document = Document(reference)
    _find_body_slot(document, REPORT_TITLE_PLACEHOLDER)
    _find_body_slot(document, REPORT_CONTENT_PLACEHOLDER)
    output.parent.mkdir(parents=True, exist_ok=True)
    if reference.resolve() != output.resolve():
        shutil.copyfile(reference, output)
